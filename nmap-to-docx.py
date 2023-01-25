#!/usr/bin/env python3
# Simple ports nmap xml report parser -> docx
# by @rextco

# pip3 install python-docx
# pip3 install python-nmap
import docx
import nmap
import os
import sys


# Input: list of xml nmap scan
# Output: array of [ip -> ports]
def parse_ports(xml_group):
	data = []

	for my_scan in xml_group:
		nm = nmap.PortScanner()
		nm.analyse_nmap_xml_scan(open(my_scan).read())

		for host in nm.all_hosts():
			parse_host = nm[host]

			my_ip, my_ports, my_detail = '', [], []
			if parse_host.state() == 'up':
				my_ip = parse_host['addresses']['ipv4']

				# TCP
				if len(parse_host.all_tcp()):
					for port in parse_host.all_tcp():
						my_ports.append(f'{port}/tcp')
						detail = parse_host['tcp'][port]['product'] + ' ' + parse_host['tcp'][port]['version']
						my_detail.append(detail)

				# UDP
				if len(parse_host.all_udp()):
					for port in parse_host.all_udp():
						my_ports.append(f'{port}/udp')
						detail = parse_host['udp'][port]['product'] + ' ' + parse_host['udp'][port]['version']
						my_detail.append(detail)

			data.append([my_ip, my_ports, my_detail])
	return data


# Input: servers.txt
# Output: List of files [xml nmap scan]
def traverse_folder(file):
	xml = []

	handler = open(file, 'rt')
	for ip in handler:
		# TCP
		my_xml = f'nmap/{ip.strip()}/recon_targeted.xml'
		if os.path.exists(my_xml):
			xml.append(my_xml)

		# UDP
		my_xml = f'nmap/{ip.strip()}/recon_targeted_udp.xml'
		if os.path.exists(my_xml):
			xml.append(my_xml)

	handler.close()
	return xml


# Input: array of [ip -> ports]
# Output: table on docx
def make_report(my_data):
	doc = docx.Document()

	for i in range(0, len(my_data)):
		my_ip = my_data[i][0]
		my_ports = my_data[i][1]
		my_detail = my_data[i][2]

		# sub-title
		doc.add_heading(f'Anexo - Puertos y servicios expuestos en {my_ip}', 1)
		# https://github.com/python-openxml/python-docx/blob/da75fcf01f7f322e846e2ac3e1936aedd766acc8/docx/enum/style.py
		table = doc.add_table(rows=1, cols=3, style='Light List Accent 1')

		# Table header
		header = table.rows[0].cells
		header[0].text = 'IP'
		header[1].text = 'Puerto'
		header[2].text = 'Detalle'

		# Table ports rows
		for counter in range(0, len(my_ports)):
			row = table.add_row().cells
			row[0].text = my_ip if counter == 0 else '' 	# Only print IP on first row
			row[1].text = my_ports[counter]
			row[2].text = my_detail[counter]

	doc.save('report.docx')


if __name__ == '__main__':
	xml = traverse_folder('servers.txt')
	ports_by_ip = parse_ports(xml)
	# ports_by_ip = parse_ports(['nmap/10.10.110.15/recon_targeted_udp.xml']) 	# for test or whatever
	make_report(ports_by_ip)
